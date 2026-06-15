from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "天玑个人基金组合分析与回测系统 V1.0_源程序鉴别材料.docx"
APP_NAME = "天玑个人基金组合分析与回测系统 V1.0"
LINES_PER_PAGE = 50
PAGES_PER_PART = 30

# Declared complete original-source order used before extracting the continuous
# first and last 30 pages. It follows startup -> GUI -> application flow ->
# domain services -> infrastructure. Third-party code and document generators
# are intentionally excluded.
SOURCE_ORDER = [
    "main.py",
    "app_info.py",
    "app_logging.py",
    "ui_app.py",
    "run_portfolio.py",
    "portfolio_analysis.py",
    "quant_service.py",
    "position_service.py",
    "portfolio_store.py",
    "fund_registry.py",
    "data_provider.py",
    "indicators/__init__.py",
    "indicators/fund_indicators.py",
    "indicators/portfolio_indicators.py",
    "backtest/__init__.py",
    "backtest/engine.py",
    "recommendation/__init__.py",
    "recommendation/models.py",
    "recommendation/scoring.py",
    "recommendation/engine.py",
    "recommendation/service.py",
]


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def load_source_lines() -> list[tuple[str, int, str]]:
    lines: list[tuple[str, int, str]] = []
    for relative in SOURCE_ORDER:
        path = ROOT / relative
        raw_lines = path.read_text(encoding="utf-8-sig").splitlines()
        for line_number, text in enumerate(raw_lines, start=1):
            # Preserve every physical source line, including original blank
            # lines, comments and indentation.
            lines.append((relative.replace("\\", "/"), line_number, text.expandtabs(4)))
    return lines


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.35)
    section.header_distance = Cm(0.55)
    section.footer_distance = Cm(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    normal.font.size = Pt(7.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{APP_NAME}  源程序鉴别材料")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8.5)
    add_page_field(footer)
    run = footer.add_run(" 页")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(8.5)


def add_source_page(
    doc: Document,
    page_lines: list[tuple[str, int, str]],
    material_line_start: int,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(10.0)
    paragraph.paragraph_format.keep_together = True

    current_file = None
    for offset, (relative, original_line, text) in enumerate(page_lines):
        if current_file != relative:
            current_file = relative
        material_line = material_line_start + offset
        prefix = f"{material_line:04d} "
        display_text = text if text else " "
        run = paragraph.add_run(prefix + display_text)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        run.font.size = Pt(7.5)
        if offset < len(page_lines) - 1:
            run.add_break()


def main() -> None:
    all_lines = load_source_lines()
    part_size = LINES_PER_PAGE * PAGES_PER_PART
    if len(all_lines) <= part_size * 2:
        selected = all_lines
    else:
        selected = all_lines[:part_size] + all_lines[-part_size:]

    if len(selected) != part_size * 2:
        raise RuntimeError(
            f"Expected {part_size * 2} selected lines for 60 pages, got {len(selected)}"
        )

    doc = Document()
    configure(doc)
    for page_index in range(60):
        start = page_index * LINES_PER_PAGE
        page_lines = selected[start : start + LINES_PER_PAGE]
        add_source_page(doc, page_lines, start + 1)
        if page_index < 59:
            doc.add_page_break()

    properties = doc.core_properties
    properties.title = f"{APP_NAME} 源程序鉴别材料"
    properties.subject = "计算机软件著作权登记普通交存源程序鉴别材料"
    properties.comments = (
        "完整原创源码按既定顺序排列并保留空行、注释和缩进后，"
        "截取前连续30页和后连续30页；每页50个原始物理行，共60页。"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)

    front_end = all_lines[part_size - 1]
    back_start = all_lines[-part_size]
    print(f"output={OUTPUT}")
    print(f"all_physical_lines={len(all_lines)}")
    print(f"selected_lines={len(selected)}")
    print(f"pages={len(selected) // LINES_PER_PAGE}")
    print(f"front_end={front_end[0]}:{front_end[1]}")
    print(f"back_start={back_start[0]}:{back_start[1]}")


if __name__ == "__main__":
    main()

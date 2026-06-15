from __future__ import annotations

from datetime import date
from pathlib import Path
import tempfile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SCREENSHOT = (
    Path(tempfile.gettempdir())
    / "codex-clipboard-474705f2-0aaa-455f-bf9e-884a4c985243.png"
)
OUTPUT = ROOT / "output"

APP_NAME = "天玑个人基金组合分析与回测系统 V1.0"
APP_SHORT_NAME = "天玑基金组合分析系统"
VERSION = "V1.0"
RELEASE_DATE = "2026年6月12日"
DISCLAIMER = (
    "本软件仅用于个人基金数据分析、学习记录、组合分析和回测辅助，不提供自动交易功能。"
    "所有收益、风险、回测和策略建议均基于历史数据、公开数据及人工规则，仅供研究参考，"
    "不构成任何投资建议、收益承诺或交易指令。历史数据和回测结果不代表未来表现，"
    "使用者应自行核对数据、独立作出判断并承担使用风险。"
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="FFFFFF")
        set_cell_shading(table.rows[0].cells[idx], "1F4E78")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if len(table.rows) % 2 == 1:
                set_cell_shading(cells[idx], "EAF2F8")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中右键更新域"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, placeholder, end])


def setup_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 24, "17365D"),
        ("Heading 1", 16, "17365D"),
        ("Heading 2", 13, "1F4E78"),
        ("Heading 3", 11, "2F5597"),
    ]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    if "Caption CN" not in [s.name for s in styles]:
        caption = styles.add_style("Caption CN", WD_STYLE_TYPE.PARAGRAPH)
        caption.font.name = "宋体"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        caption.font.size = Pt(9)
        caption.font.color.rgb = RGBColor(89, 89, 89)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = f"{APP_NAME}  |  {title}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"{APP_SHORT_NAME}  {VERSION}    第 ")
    add_field(footer, "PAGE")
    footer.add_run(" 页")
    for run in footer.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)


def add_cover(doc: Document, doc_title: str, subtitle: str) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(APP_NAME)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(25)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 54, 93)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_title)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)
    for _ in range(7):
        doc.add_paragraph()
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["软件全称", APP_NAME],
            ["软件简称", APP_SHORT_NAME],
            ["版本号", VERSION],
            ["文档状态", "V1.0 正式申请准备稿"],
            ["编制日期", RELEASE_DATE],
        ],
        [4, 11],
    )
    p = doc.add_paragraph(DISCLAIMER)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(12)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)
    doc.add_page_break()


def add_revision_and_toc(doc: Document, purpose: str) -> None:
    doc.add_heading("文档说明", level=1)
    doc.add_paragraph(purpose)
    add_table(
        doc,
        ["版本", "日期", "状态", "说明"],
        [[VERSION, "2026-06-12", "正式申请准备稿", "依据冻结版功能基线、README 与当前程序整理"]],
        [2, 3, 3, 8],
    )
    doc.add_heading("目录", level=1)
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_paragraph("提示：在 Microsoft Word 中打开文档后，可右键目录并选择“更新域”生成页码。")
    doc.add_page_break()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_picture(doc: Document, path: Path, caption: str, width_cm: float = 16.0) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(width_cm))
        doc.add_paragraph(caption, style="Caption CN")
    else:
        p = doc.add_paragraph(f"[图片未找到：{path.name}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    run.font.size = Pt(9)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)


def build_user_manual() -> Path:
    doc = Document()
    setup_document(doc, "用户操作手册")
    add_cover(doc, "用户操作手册", "运行、数据管理、分析回测与报告输出说明")
    add_revision_and_toc(
        doc,
        "本文档面向软件使用者，说明安装启动、页面功能、数据录入、分析回测、报告导出、"
        "日志排查和风险提示。示例采用仓库随附的 V1.0 模拟演示持仓，不含真实个人资产信息。",
    )

    doc.add_heading("1 软件简介", level=1)
    doc.add_paragraph(
        f"{APP_NAME}是一套 Python 桌面端个人基金组合分析与回测辅助软件。"
        "软件围绕基金池、真实净值缓存、持仓批次、统一估值、收益风险指标、组合分析、"
        "定投与均线回测、规则化策略建议和多格式报告输出形成完整操作链路。"
    )
    doc.add_heading("1.1 适用范围", level=2)
    add_bullets(
        doc,
        [
            "记录和维护个人基金池、持仓、交易流水及买入批次。",
            "读取公开基金净值，结合本地真实缓存进行估值和历史分析。",
            "进行收益风险分析、组合级分析、定投模拟和规则策略回测。",
            "生成 CSV、PNG 和 Markdown 报告，供个人学习、复核和留档。",
            "不连接证券或基金交易账户，不执行自动申购、赎回或交易。",
        ],
    )
    doc.add_heading("1.2 标准演示数据", level=2)
    doc.add_paragraph(
        "仓库当前随附的 portfolio_data.json 标记为 data_profile=demo，是 V1.0 标准模拟演示持仓。"
        "演示组合包含 7 只基金，覆盖 A 股宽基、黄金、短债、科技成长和美股纳斯达克等类别。"
        "GUI 页眉显示“V1.0 模拟演示数据”，用于截图、操作说明和测试，不代表真实资产。"
    )
    add_table(
        doc,
        ["类别", "演示内容"],
        [
            ["A股宽基", "富国中证A500ETF联接A、南方深证成份ETF联接A"],
            ["黄金", "易方达黄金ETF联接A"],
            ["短债", "长城短债债券A、嘉实短债债券A"],
            ["科技成长", "华夏上证科创板综合ETF联接A"],
            ["美股纳斯达克", "华安纳斯达克100ETF联接(QDII)A"],
        ],
    )

    doc.add_heading("2 运行环境与安装", level=1)
    doc.add_heading("2.1 推荐环境", level=2)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ["操作系统", "Windows 10/11（V1.0 主要适配环境）"],
            ["Python", "Python 3.10 或更高版本"],
            ["网络", "首次获取或刷新基金净值时需要访问第三方公开数据源"],
            ["显示", "建议 1920×1080 或更高分辨率；低分辨率可使用页面滚动区域"],
            ["主要依赖", "xalpha、pandas、numpy、matplotlib；Tkinter/ttk 随 Python 提供"],
        ],
    )
    doc.add_heading("2.2 安装步骤", level=2)
    add_steps(
        doc,
        [
            "进入项目目录 xalpha_portfolio_analyzer。",
            "创建并激活 Python 虚拟环境。",
            "执行 pip install -r requirements.txt 安装依赖。",
            "确认项目目录具有读取和写入 data、logs、output 等目录的权限。",
        ],
    )
    add_code_block(
        doc,
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate\n"
        "pip install -r requirements.txt",
    )
    doc.add_heading("2.3 启动方式", level=2)
    doc.add_paragraph("V1.0 统一 GUI 启动入口：")
    add_code_block(doc, "python main.py")
    doc.add_paragraph("命令行完整分析入口：")
    add_code_block(doc, "python run_portfolio.py")
    doc.add_paragraph(
        "原 python ui_app.py 入口为兼容入口。正式操作、截图和材料说明统一使用 python main.py。"
    )

    doc.add_heading("3 主界面与通用操作", level=1)
    doc.add_paragraph(
        "软件启动后进入投资组合仪表盘。顶部显示软件全称、版本、数据状态、输出与报告快捷按钮；"
        "中部以标签页组织 12 个功能页面。点击“重新分析”会在后台执行数据更新、估值、指标、"
        "回测、建议和报告流程，界面保持可响应，并在完成后刷新页面。"
    )
    add_picture(doc, SCREENSHOT, "图 3-1 总览与图表页面（V1.0 模拟演示数据）", 17.0)
    add_table(
        doc,
        ["页面", "用途"],
        [
            ["总览与图表", "查看组合总资产、收益、配置和核心图表"],
            ["持仓明细", "查看最新净值、市值、成本、收益、权重和趋势"],
            ["交易流水", "登记确认买入/卖出，查看和撤销允许撤销的记录"],
            ["持仓批次", "管理多笔买入批次、剩余份额和成本"],
            ["基金池管理", "维护基金代码、名称、分类、类型、启停和定投配置"],
            ["数据中心", "查看净值日期、数据来源、缓存状态和错误"],
            ["收益与风险", "查看阶段收益和单基金风险指标"],
            ["定投模拟", "查看预设基金的基础定投结果"],
            ["量化指标", "查看基金与组合指标数据集"],
            ["策略回测", "设置参数并比较固定定投、动态定投和均线策略"],
            ["策略建议", "查看规则化动作、评分、金额、理由和风险信号"],
            ["运行日志", "查看刷新、估值、分析步骤和异常信息"],
        ],
    )

    doc.add_heading("4 基金池管理", level=1)
    doc.add_heading("4.1 添加基金", level=2)
    add_steps(
        doc,
        [
            "进入“基金池管理”页面，点击添加。",
            "填写唯一基金代码、基金名称、资产分类和基金类型。",
            "选择数据源 xalpha，并设置是否启用。",
            "若为定投基金，填写频率、基础金额、是否允许暂停/增强和最大增强倍数。",
            "保存后执行重新分析，检查数据中心中的净值日期和状态。",
        ],
    )
    doc.add_heading("4.2 编辑、启用和停用", level=2)
    doc.add_paragraph(
        "基金代码是跨模块关联主键，编辑基金时不可修改代码。名称、分类、类型和定投参数可调整。"
        "停用后基金不参与当前分析流程，但配置仍保留；重新启用后恢复参与。"
    )
    doc.add_heading("4.3 删除基金", level=2)
    doc.add_paragraph(
        "删除操作会明确提示，并同步处理该基金的持仓、交易和批次记录。删除前应确认已完成必要备份。"
    )

    doc.add_heading("5 数据更新与缓存", level=1)
    doc.add_heading("5.1 重新分析", level=2)
    add_steps(
        doc,
        [
            "保持网络连接，点击顶部“重新分析”。",
            "程序优先尝试从远程公开数据源更新基金净值。",
            "成功数据写入 data_xalpha/cache/<基金代码>.csv。",
            "同一轮后续估值、指标和报告复用本轮缓存，避免重复请求。",
            "分析完成后查看顶部状态、数据中心和运行日志。",
        ],
    )
    doc.add_heading("5.2 数据状态判断", level=2)
    add_table(
        doc,
        ["状态", "含义", "建议"],
        [
            ["远程更新成功", "取得新的正式净值并写入缓存", "核对净值截止日期后继续分析"],
            ["使用本地缓存", "远程失败或无新增记录，使用此前真实净值", "检查缓存日期，不要把缓存理解为当日已更新"],
            ["数据过期", "净值日期明显早于合理披露日期", "检查网络、数据源和基金公告"],
            ["数据失败", "无远程数据且无可用缓存", "查看日志错误，相关指标会显示数据不足"],
        ],
    )
    doc.add_heading("5.3 QDII 与披露滞后", level=2)
    doc.add_paragraph(
        "QDII 基金受境外市场交易日、时差和估值流程影响，正式净值通常比境内基金滞后一个或多个交易日。"
        "软件按各基金最新正式净值估算当前持仓；组合历史指标则截止到所有有效成分共同拥有的最新正式净值日，"
        "避免把尚未公布的 QDII 收益错误地当作 0%。详情页补全只使用已正式公布的单位净值，不采用盘中估算值。"
    )

    doc.add_heading("6 持仓、交易与批次", level=1)
    doc.add_heading("6.1 录入买入批次", level=2)
    add_steps(
        doc,
        [
            "进入“持仓批次”或“交易流水”页面。",
            "选择基金代码，填写确认日期、确认金额、确认净值、手续费和备注。",
            "确认份额可由系统按确认金额÷确认净值计算；若手工份额与金额冲突，以确认金额为主重算。",
            "保存后系统按剩余批次重建基金持仓，并在持仓明细中显示估值。",
        ],
    )
    doc.add_heading("6.2 卖出与 FIFO", level=2)
    doc.add_paragraph(
        "卖出交易按先进先出（FIFO）消耗历史买入批次。V1.0 不支持直接撤销已影响 FIFO 批次的卖出，"
        "需要通过更正记录处理。录入前应核对基金、确认日期、金额、净值、份额和手续费。"
    )
    doc.add_heading("6.3 历史持仓估算", level=2)
    doc.add_paragraph(
        "缺少真实确认份额的旧持仓可通过“历史持仓估算录入”迁移。系统会根据已有市值和净值推算份额，"
        "并显示“估算”或“待估值”标记。该方式仅用于过渡，精确分析应补录真实批次。"
    )
    doc.add_heading("6.4 估值公式", level=2)
    add_code_block(
        doc,
        "确认份额 = 确认金额 / 确认净值\n"
        "当前市值 = 剩余确认份额 × 最新单位净值\n"
        "当前收益 = 当前市值 - 持仓成本\n"
        "当前收益率 = 当前收益 / 持仓成本",
    )

    doc.add_heading("7 收益、风险与组合分析", level=1)
    doc.add_paragraph(
        "完成重新分析后，可在“收益与风险”和“量化指标”页面查看阶段收益、年化收益、年化波动率、"
        "最大回撤、夏普比率、卡玛比率、RSI 和移动平均线等结果。阶段收益目标日期无净值时，"
        "采用该日期之前最近的正式交易日记录。"
    )
    add_picture(doc, OUTPUT / "asset_allocation.png", "图 7-1 模拟演示组合资产配置图", 13.5)
    add_picture(doc, OUTPUT / "category_return.png", "图 7-2 模拟演示组合类别近 30 日平均收益", 13.5)
    add_picture(doc, OUTPUT / "risk_comparison.png", "图 7-3 模拟演示基金风险收益矩阵", 13.5)
    doc.add_paragraph(
        "组合净值按分析时的当前持仓权重合成，不还原历史每个时点的真实动态权重。"
        "组合截止日期可能早于部分境内基金的最新单位净值日期，这是披露滞后成分共同日期对齐的结果。"
    )

    doc.add_heading("8 定投模拟与策略回测", level=1)
    doc.add_heading("8.1 定投模拟", level=2)
    doc.add_paragraph(
        "“定投模拟”页面展示预设基金的基础每日定投演示。按月定投曲线出现阶梯状直线属于正常现象，"
        "因为资金仅在计划日期增加。"
    )
    doc.add_heading("8.2 策略回测操作", level=2)
    add_steps(
        doc,
        [
            "进入“策略回测”页面并选择基金。",
            "设置起止日期、每日/每周/每月频率和基础金额。",
            "运行回测，比较固定定投、MA60 动态定投和 MA20/MA60 均线策略。",
            "查看累计投入、期末资产、收益、回撤、年化指标、交易次数和资金曲线。",
            "需要留档时打开 output 目录查看 backtest_result.csv 和 backtest_curve.csv。",
        ],
    )
    doc.add_heading("8.3 结果解释", level=2)
    doc.add_paragraph(
        "动态定投在净值低于 MA60 时投入 1.5 倍基础金额，高于或等于 MA60 时投入 0.5 倍。"
        "均线策略在 MA20 上穿并高于 MA60 且空仓时买入，在 MA20 不高于 MA60 且持仓时卖出。"
        "回测未完整模拟费用、滑点、限购、暂停申购和到账延迟，不能等同真实交易收益。"
    )

    doc.add_heading("9 策略建议", level=1)
    doc.add_paragraph(
        "策略建议使用趋势、位置、风险、仓位和定投配置生成可解释规则结果。"
        "页面展示当前动作、综合评分、定投倍率与金额、卖出风险、数据日期、数据来源、置信度和理由。"
    )
    add_table(
        doc,
        ["信息", "说明"],
        [
            ["当前建议", "重点关注、小额买入、持有、观察、减仓、暂停定投等规则动作"],
            ["建议金额", "定投基础金额×建议倍率；非定投基金通常为 0"],
            ["卖出信号", "止盈、风险减仓或止损风险提示，不直接执行交易"],
            ["综合评分", "趋势、位置、风险和持仓四维加权结果"],
            ["置信度", "表示数据和规则完整度，不表示建议一定正确"],
            ["数据质量", "反映净值来源、日期、缓存和配置完整性"],
        ],
    )
    doc.add_paragraph(
        "QDII 建议会附带净值滞后提示；货币基金不使用均线和 RSI 择时；债券基金采用更保守的风险阈值。"
    )

    doc.add_heading("10 报告导出与文件位置", level=1)
    doc.add_heading("10.1 快捷操作", level=2)
    add_bullets(
        doc,
        [
            "“打开输出”：打开 output 目录。",
            "“查看报告”：打开最近一次组合分析 Markdown 报告。",
            "完成分析后，确认状态为全部完成或查看部分失败步骤。",
        ],
    )
    add_table(
        doc,
        ["文件/目录", "内容"],
        [
            ["output/summary.csv", "组合总资产、成本、收益和更新时间"],
            ["output/holdings_result.csv", "统一估值后的持仓明细"],
            ["output/period_return.csv", "阶段收益和净值截止日期"],
            ["output/risk_report.csv", "收益风险指标"],
            ["output/indicator_dataset.csv", "基金与组合指标数据集"],
            ["output/portfolio_nav.csv", "组合净值曲线"],
            ["output/backtest_result.csv", "策略回测绩效"],
            ["output/backtest_curve.csv", "策略资金曲线"],
            ["output/*.png", "资产配置、类别收益、风险比较图"],
            ["output/fund_recommendation_report.md", "策略建议报告"],
            ["portfolio_report.md", "最近一次组合分析报告"],
            ["data/history/fund_recommendations.csv", "历史建议记录"],
        ],
    )
    doc.add_paragraph("CSV 使用 UTF-8 with BOM，便于 Windows Excel 正确识别中文。")

    doc.add_heading("11 日志与故障排查", level=1)
    doc.add_paragraph("运行日志位于 logs/xalpha_portfolio.log，也可在“运行日志”页面查看。")
    add_table(
        doc,
        ["现象", "可能原因", "处理方法"],
        [
            ["点击重新分析后日期未变化", "上游尚未公布新净值、远程失败或基金本身披露滞后", "查看数据中心来源、净值日期和日志，不以当前自然日判断基金已更新"],
            ["QDII 日期比境内基金早", "境外市场、时差和估值披露机制", "属于常见情况；关注正式净值日期和组合共同截止日"],
            ["部分指标显示数据不足", "历史长度不足、缓存缺失或更新失败", "检查基金是否启用、缓存文件和日志错误"],
            ["持仓收益与手工金额不同", "软件按剩余份额×最新净值统一估值", "核对批次份额、成本、手续费和最新净值"],
            ["分析提示部分完成", "某一输出步骤失败，其余步骤已保存", "在日志中查找失败步骤，不把旧输出当作本轮结果"],
            ["中文 CSV 打开乱码", "表格软件未按 UTF-8 识别", "使用 Excel 导入或选择 UTF-8 编码"],
        ],
    )

    doc.add_heading("12 数据备份与迁移", level=1)
    add_bullets(
        doc,
        [
            "核心业务数据文件为 fund_pool.json 和 portfolio_data.json。",
            "切换整套持仓时主要替换 portfolio_data.json；基金范围不同还需同步 fund_pool.json。",
            "净值缓存位于 data_xalpha/cache，可重新获取，不等同核心持仓数据。",
            "旧版 portfolio_data.json 迁移前会生成 portfolio_data.backup_YYYYMMDD_HHMMSS.json。",
            "替换或编辑业务数据前应关闭软件并保留原文件备份。",
        ],
    )

    doc.add_heading("13 常见问题", level=1)
    add_table(
        doc,
        ["问题", "回答"],
        [
            ["本软件会自动交易吗？", "不会。软件不连接交易账户，也不执行申购、赎回或买卖。"],
            ["本地缓存是模拟数据吗？", "不是。缓存是此前从公开数据源成功取得的真实历史净值。"],
            ["能否直接使用当前持仓作为演示数据？", "可以。当前 portfolio_data.json 已标记为模拟演示数据。"],
            ["为什么组合指标日期比单只基金晚/早？", "组合取所有有效成分共同拥有的最新正式净值日。"],
            ["策略建议是否代表投资建议？", "不代表。它是固定规则模型生成的研究参考。"],
            ["删除 portfolio_data.json 会怎样？", "程序创建空白用户数据结构，不会自动再写入另一套演示金额。"],
        ],
    )

    doc.add_heading("14 已知限制与免责声明", level=1)
    add_bullets(
        doc,
        [
            "xalpha 及上游公开数据源可能因网络、接口或页面变化而更新失败。",
            "远程失败时使用最近一次成功缓存，缓存不代表当日已经更新。",
            "QDII 正式净值通常比境内基金滞后一个或多个交易日。",
            "历史持仓缺少真实确认份额时只能估算，并显示数据质量标签。",
            "组合历史净值按当前持仓权重合成，不还原历史动态持仓。",
            "回测未完整模拟费用、滑点、限购、暂停申购、到账延迟和全部交易约束。",
            "卖出采用 FIFO，V1.0 不能直接撤销卖出交易。",
            "GUI 主要面向 Windows 10/11，V1.0 尚无独立安装包、完整自动化测试和持续集成。",
        ],
    )
    p = doc.add_paragraph(DISCLAIMER)
    p.paragraph_format.space_before = Pt(10)
    for run in p.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)

    doc.add_heading("附录 A 标准演示操作路径", level=1)
    add_steps(
        doc,
        [
            "执行 python main.py 启动软件，确认页眉显示软件全称、V1.0 和模拟演示数据标识。",
            "在基金池管理页检查 7 只演示基金及分类、类型、定投配置。",
            "点击重新分析，在数据中心检查各基金净值日期、数据来源和缓存状态。",
            "在持仓批次页检查模拟买入批次，在持仓明细页核对统一估值。",
            "依次查看收益与风险、量化指标、定投模拟和策略回测。",
            "查看策略建议的动作、金额、评分、理由、数据质量和 QDII 提示。",
            "打开 output 目录，核对 CSV、PNG、Markdown 报告和运行日志。",
        ],
    )

    path = DOCS / f"{APP_NAME}_用户操作手册.docx"
    doc.save(path)
    return path


def build_design_spec() -> Path:
    doc = Document()
    setup_document(doc, "软件设计说明书")
    add_cover(doc, "软件设计说明书", "总体架构、数据模型、核心算法与模块边界")
    add_revision_and_toc(
        doc,
        "本文档面向软件设计审查和软著材料准备，说明 V1.0 的总体架构、模块职责、数据模型、"
        "核心处理链路、第三方依赖、原创业务边界和已知限制。",
    )

    doc.add_heading("1 引言", level=1)
    doc.add_heading("1.1 编写目的", level=2)
    doc.add_paragraph(
        "本文档用于描述天玑个人基金组合分析与回测系统 V1.0 的技术结构和实现边界，"
        "为软件维护、功能核查、测试、用户文档编制和计算机软件著作权申请提供一致依据。"
    )
    doc.add_heading("1.2 软件定位", level=2)
    doc.add_paragraph(
        "系统定位为个人基金数据分析、学习记录、组合分析与回测辅助工具。"
        "系统不连接交易账户，不执行自动交易，不承诺收益，不构成投资建议。"
    )
    doc.add_heading("1.3 设计原则", level=2)
    add_bullets(
        doc,
        [
            "以基金代码作为跨模块稳定标识，避免同名或改名造成关联错误。",
            "使用真实公开净值和真实缓存，不在数据失败时伪造随机结果。",
            "持仓估值采用统一服务，保证 GUI、CSV、报告和建议口径一致。",
            "数据质量、缓存状态和净值截止日期必须可见。",
            "已知限制通过明确标识和文档说明控制，不在 V1.0 无边界扩展功能。",
        ],
    )

    doc.add_heading("2 总体架构", level=1)
    doc.add_paragraph(
        "系统采用本地桌面应用架构。Tkinter GUI 和命令行入口调用应用层编排，"
        "应用层协调基金池、数据源缓存、持仓存储、统一估值、指标、回测、建议和报告模块。"
        "业务数据主要以 JSON/CSV 文件保存，图表和报告输出到本地目录。"
    )
    add_table(
        doc,
        ["层次", "模块", "职责"],
        [
            ["入口与展示层", "main.py、ui_app.py、run_portfolio.py", "GUI 启动、页面交互、命令行分析和流程编排"],
            ["应用服务层", "position_service.py、quant_service.py、recommendation/service.py", "统一估值、指标回测编排、建议批处理"],
            ["领域逻辑层", "portfolio_store.py、indicators/、backtest/、recommendation/", "持仓交易、指标、策略和评分规则"],
            ["基础设施层", "fund_registry.py、data_provider.py、app_logging.py", "配置、数据源、缓存、日志和异常状态"],
            ["数据与输出层", "JSON、CSV、PNG、Markdown", "业务数据持久化、缓存和分析结果"],
        ],
    )
    doc.add_heading("2.1 模块依赖", level=2)
    add_code_block(
        doc,
        "main.py\n"
        "└─ ui_app.py\n"
        "   ├─ app_info.py\n"
        "   ├─ fund_registry.py\n"
        "   ├─ data_provider.py\n"
        "   ├─ portfolio_store.py\n"
        "   ├─ quant_service.py\n"
        "   ├─ recommendation/\n"
        "   └─ run_portfolio.py\n"
        "      └─ portfolio_analysis.py\n"
        "\n"
        "quant_service.py → indicators/ + backtest/\n"
        "data_provider.py → xalpha + data_xalpha/cache/",
    )
    doc.add_heading("2.2 运行部署", level=2)
    add_bullets(
        doc,
        [
            "主要运行平台：Windows 10/11。",
            "运行时：Python 3.10 及以上。",
            "部署形态：源代码与依赖环境本地运行，V1.0 未提供独立安装包。",
            "GUI 统一入口：python main.py；CLI 入口：python run_portfolio.py。",
        ],
    )

    doc.add_heading("3 模块设计", level=1)
    modules = [
        ("3.1 应用标识模块 app_info.py", "集中定义软件全称、简称、版本号和统一免责声明，供 GUI、CLI 和报告复用。"),
        ("3.2 GUI 模块 ui_app.py", "实现 12 个标签页、后台分析线程、进度状态、内嵌图表、表格展示、文件打开和用户输入校验。"),
        ("3.3 流程编排 run_portfolio.py", "按步骤执行净值准备、统一估值、收益风险、量化指标、回测、建议和报告输出，并区分全部完成、部分完成和失败。"),
        ("3.4 基金池 fund_registry.py", "读取和校验 fund_pool.json，支持基金增删改、启停、类型、分类和定投配置维护。"),
        ("3.5 数据源 data_provider.py", "封装 xalpha 普通/货币基金接口、正式单位净值补全、缓存读写、远程失败回退和 provider_status 状态。"),
        ("3.6 持仓存储 portfolio_store.py", "维护 holdings、transactions、lots、position_summary，处理买卖交易、FIFO、数据迁移和备份。"),
        ("3.7 统一估值 position_service.py", "根据剩余份额和最新正式单位净值计算基金级及组合级市值、成本、收益、收益率和权重。"),
        ("3.8 指标 indicators/", "fund_indicators.py 计算单基金通用指标；portfolio_indicators.py 对齐成分日期并计算组合净值和组合风险指标。"),
        ("3.9 回测 backtest/", "engine.py 实现固定定投、MA60 动态定投和 MA20/MA60 均线策略，以及资金曲线和绩效统计。"),
        ("3.10 量化服务 quant_service.py", "组织指标数据集、组合净值、回测结果与资金曲线，并输出应用层 CSV。"),
        ("3.11 策略建议 recommendation/", "models 定义结构；scoring 计算四维分；engine 生成普通/定投/卖出规则；service 批量生成、保存历史和报告。"),
        ("3.12 分析报告 portfolio_analysis.py", "计算阶段收益、风险、基础定投，生成资产配置、类别收益、风险矩阵和组合 Markdown 报告。"),
        ("3.13 日志 app_logging.py", "配置轮转日志，记录数据源、缓存、净值日期、估值方式、步骤状态和异常原因。"),
    ]
    for heading, body in modules:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)

    doc.add_heading("4 数据模型", level=1)
    doc.add_heading("4.1 基金池模型", level=2)
    add_table(
        doc,
        ["字段", "类型/示例", "说明"],
        [
            ["code", "字符串 007194", "基金唯一代码，跨模块关联主键"],
            ["name", "字符串", "基金显示名称，要求不重复"],
            ["category", "字符串 短债", "资产分类，用于配置和类别收益"],
            ["fund_type", "normal/money/qdii 等", "数据源和特殊规则依据"],
            ["enabled", "布尔值", "是否参与当前分析"],
            ["data_source", "xalpha", "数据源标识"],
            ["is_dca", "布尔值", "是否启用定投建议"],
            ["dca_frequency", "monthly 等", "定投频率"],
            ["dca_base_amount", "数值", "定投基础金额"],
            ["dca_allow_pause/increase", "布尔值", "允许暂停或增强"],
            ["dca_max_multiplier", "数值", "最大增强倍率"],
        ],
    )
    doc.add_heading("4.2 组合数据模型", level=2)
    add_table(
        doc,
        ["对象", "主要字段", "用途"],
        [
            ["holdings", "code、cost、shares、quality 等", "基金级汇总持仓和迁移状态"],
            ["transactions", "id、code、side、date、amount、nav、shares、fee", "确认买入/卖出流水"],
            ["lots", "lot_id、code、buy_date、confirmed_shares、remaining_shares、cost", "逐笔买入批次和 FIFO 剩余份额"],
            ["position_summary", "market_value、cost、profit、profit_rate、updated_at", "最近一次组合统一估值摘要"],
            ["data_profile", "demo/user", "区分模拟演示数据与用户数据"],
            ["version", "当前为 3", "数据迁移版本控制"],
        ],
    )
    doc.add_heading("4.3 净值缓存模型", level=2)
    doc.add_paragraph(
        "每只基金缓存为 data_xalpha/cache/<code>.csv，主要包含日期、单位净值 netvalue、"
        "累计净值 totvalue 及数据来源信息。provider_status.json 记录最近一次远程状态、"
        "缓存状态、净值日期、是否过期和错误原因。"
    )
    doc.add_heading("4.4 输出模型", level=2)
    add_table(
        doc,
        ["类型", "格式", "主要文件"],
        [
            ["组合与持仓", "CSV", "summary.csv、holdings_result.csv"],
            ["收益风险", "CSV", "period_return.csv、risk_report.csv、indicator_dataset.csv"],
            ["组合与回测曲线", "CSV", "portfolio_nav.csv、backtest_result.csv、backtest_curve.csv"],
            ["图表", "PNG", "asset_allocation.png、category_return.png、risk_comparison.png"],
            ["报告", "Markdown", "portfolio_report.md、fund_recommendation_report.md"],
            ["历史记录", "CSV", "data/history/fund_recommendations.csv"],
        ],
    )

    doc.add_heading("5 核心数据链路", level=1)
    doc.add_heading("5.1 完整分析流程", level=2)
    add_steps(
        doc,
        [
            "读取 fund_pool.json，取得启用基金和配置。",
            "对每只启用基金请求 xalpha 远程数据，必要时补入已正式公布单位净值。",
            "成功数据写入真实缓存；远程失败时读取已有缓存并记录状态。",
            "读取 portfolio_data.json 的持仓、交易和批次。",
            "统一估值服务按剩余份额和最新净值生成持仓行与组合摘要。",
            "收益风险模块计算阶段收益、单基金指标和图表。",
            "量化服务计算基金指标数据集、组合净值与策略回测。",
            "建议服务结合指标、持仓、基金类型和定投配置生成建议。",
            "GUI、CSV、PNG、Markdown 和日志消费同一轮分析结果。",
        ],
    )
    doc.add_heading("5.2 失败与回退", level=2)
    doc.add_paragraph(
        "远程数据失败时优先使用已有有效缓存；无缓存时返回明确错误或数据不足，不生成随机净值。"
        "完整分析记录每一步状态，部分步骤失败时提示“部分完成”并列出失败步骤，避免将旧输出误认为本轮结果。"
    )

    doc.add_heading("6 统一估值设计", level=1)
    doc.add_heading("6.1 设计目标", level=2)
    doc.add_paragraph(
        "统一估值消除总览、持仓明细、报告和建议分别读取旧金额造成的跨页面不一致。"
        "所有展示和输出以 position_service 生成的持仓行和组合摘要为准。"
    )
    doc.add_heading("6.2 计算口径", level=2)
    add_code_block(
        doc,
        "remaining_shares = Σ 各未售批次剩余确认份额\n"
        "market_value = remaining_shares × latest_official_nav\n"
        "holding_cost = Σ 未售批次确认成本及对应手续费\n"
        "profit = market_value - holding_cost\n"
        "profit_rate = profit / holding_cost\n"
        "weight = market_value / portfolio_market_value",
    )
    doc.add_heading("6.3 数据质量", level=2)
    doc.add_paragraph(
        "真实批次可精确计算剩余份额和成本。历史缺失份额的数据按迁移时市值和净值估算，并标记估算/待估值。"
        "数据质量标记向 GUI 和建议模块传递，避免估算数据被误认为精确结果。"
    )

    doc.add_heading("7 收益风险与组合指标", level=1)
    doc.add_heading("7.1 单基金指标", level=2)
    add_table(
        doc,
        ["指标", "计算说明"],
        [
            ["阶段收益", "期末累计净值/期初累计净值-1，目标日缺失时向前取最近交易日"],
            ["年化收益率", "按实际自然日跨度复合年化"],
            ["年化波动率", "日收益标准差×sqrt(252)"],
            ["最大回撤", "净值相对历史峰值的最大跌幅"],
            ["夏普比率", "(年化收益率-2%)/年化波动率"],
            ["卡玛比率", "年化收益率/最大回撤"],
            ["RSI14", "14 个交易日平均涨跌强弱"],
            ["MA20/60/120/250", "对应交易日窗口移动平均"],
            ["近252日回撤", "当前净值相对近252日最高值的位置"],
        ],
    )
    doc.add_heading("7.2 组合指标与日期对齐", level=2)
    doc.add_paragraph(
        "各基金净值先标准化为起点 1.0，再按分析时当前持仓权重合成组合净值。"
        "不同基金日期在各自有效历史区间内对齐并向前填充，但组合截止日取所有有效成分共同拥有的"
        "最新正式净值日。该设计用于处理 QDII 等披露滞后基金，避免把尚未公布的收益填充为 0%。"
    )
    add_picture(doc, OUTPUT / "risk_comparison.png", "图 7-1 风险收益指标输出示例", 14.5)

    doc.add_heading("8 回测设计", level=1)
    doc.add_heading("8.1 固定定投", level=2)
    doc.add_paragraph("在每日、每周或每月计划日期按固定金额买入，累计份额乘当日净值得到资产。")
    doc.add_heading("8.2 动态定投", level=2)
    doc.add_paragraph(
        "使用 MA60 调整投入强度：净值低于 MA60 时投入 1.5 倍基础金额；"
        "净值高于或等于 MA60 时投入 0.5 倍基础金额。"
    )
    doc.add_heading("8.3 均线策略", level=2)
    doc.add_paragraph(
        "MA20 高于 MA60 且空仓时全额买入；MA20 不高于 MA60 且持仓时全部卖出；"
        "默认初始资金 10,000 元。"
    )
    doc.add_heading("8.4 绩效输出", level=2)
    add_bullets(
        doc,
        [
            "累计投入、期末资产、收益金额和现金收益率。",
            "最大回撤、年化收益率、年化波动率、夏普和卡玛。",
            "交易次数和每日策略资金曲线。",
        ],
    )
    doc.add_paragraph(
        "回测未完整模拟申购费、赎回费、滑点、限购、暂停申购和到账延迟，"
        "其结果用于历史研究，不等同实际交易收益。"
    )

    doc.add_heading("9 策略建议设计", level=1)
    doc.add_heading("9.1 输入", level=2)
    add_bullets(
        doc,
        [
            "净值与 MA20/60/120/250、MA120 斜率、RSI14、近252日回撤。",
            "年化收益、波动、最大回撤、夏普和卡玛。",
            "持仓收益率、组合权重、持有状态和数据质量。",
            "基金类型、定投基础金额、频率、暂停/增强配置。",
        ],
    )
    doc.add_heading("9.2 四维评分", level=2)
    add_code_block(
        doc,
        "普通基金综合评分 = 趋势30% + 位置25% + 风险25% + 持仓20%\n"
        "债券基金综合评分 = 趋势15% + 位置20% + 风险45% + 持仓20%",
    )
    doc.add_heading("9.3 动作与风险信号", level=2)
    add_table(
        doc,
        ["规则", "输出"],
        [
            ["评分≥80", "重点关注/可加仓；RSI 过热时转观察"],
            ["65≤评分<80", "可小额买入"],
            ["45≤评分<65", "已持有则持有，未持有则观察"],
            ["30≤评分<45", "观察；高仓位且趋势弱时适当减仓"],
            ["评分<30", "趋势弱且有持仓时部分卖出，否则暂不建议买入"],
            ["盈利>25% 且 RSI>78", "分批止盈风险提示"],
            ["长期趋势转弱且评分<30", "风险减仓提示"],
            ["亏损>25%、趋势弱且夏普<0", "止损风险提示"],
        ],
    )
    doc.add_heading("9.4 特殊基金", level=2)
    doc.add_paragraph(
        "货币基金按流动性配置，不使用均线和 RSI 择时；债券基金采用更保守风险阈值；"
        "QDII 增加净值披露滞后提示；黄金/商品提示分散属性不保证降低全部风险。"
    )

    doc.add_heading("10 GUI 设计", level=1)
    doc.add_paragraph(
        "GUI 使用 Tkinter/ttk。主窗口包含统一页眉、分析状态、快捷按钮、进度条和 12 个标签页。"
        "耗时分析在后台线程执行，通过队列将状态和结果传回主线程，避免界面长时间无响应。"
    )
    add_picture(doc, SCREENSHOT, "图 10-1 GUI 总览页面及标签页结构", 17.0)
    add_table(
        doc,
        ["设计点", "实现说明"],
        [
            ["身份一致性", "窗口标题、页眉、报告标题均使用 app_info 中的软件全称与 V1.0"],
            ["数据可见性", "显示模拟演示标识、最近更新时间、净值日期、来源和质量"],
            ["跨页一致性", "总览和持仓明细复用统一估值结果"],
            ["错误反馈", "对网络失败、缓存回退、数据不足和部分完成提供提示"],
            ["图表", "matplotlib 生成并嵌入 Tkinter 页面"],
        ],
    )

    doc.add_heading("11 报告与日志设计", level=1)
    doc.add_heading("11.1 输出", level=2)
    doc.add_paragraph(
        "分析结果分为结构化 CSV、可视化 PNG 和可阅读 Markdown。CSV 采用 utf-8-sig，"
        "报告统一包含软件全称、版本、生成时间和免责声明。"
    )
    add_picture(doc, OUTPUT / "asset_allocation.png", "图 11-1 报告资产配置图示例", 12.5)
    doc.add_heading("11.2 日志", level=2)
    add_bullets(
        doc,
        [
            "记录基金代码、名称、净值、净值日期、来源和缓存状态。",
            "记录估值方式、数据质量和组合汇总。",
            "记录完整分析各步骤开始、成功、部分失败和异常堆栈。",
            "日志写入 logs/xalpha_portfolio.log，并在 GUI 运行日志页展示。",
        ],
    )

    doc.add_heading("12 数据迁移与备份", level=1)
    doc.add_paragraph(
        "portfolio_data.json 使用 version 字段控制结构迁移，当前版本为 3。"
        "旧数据迁移前生成时间戳备份。迁移过程保留缺失真实份额的估算标签，不静默转换为精确数据。"
        "核心业务数据为 fund_pool.json、portfolio_data.json 和建议历史；缓存、日志和 output 为可再生数据。"
    )

    doc.add_heading("13 异常处理与数据安全", level=1)
    add_table(
        doc,
        ["场景", "处理"],
        [
            ["远程网络或解析失败", "记录错误并读取有效缓存；无缓存时返回失败"],
            ["净值尚未更新", "保留最近正式净值和截止日期，不使用盘中估值替代"],
            ["组合无有效权重", "执行零有效权重保护，不生成误导性组合指标"],
            ["部分输出失败", "总体状态标为部分完成，列出未生成步骤"],
            ["旧数据缺少份额", "估算并标记质量，提示补录真实批次"],
            ["删除基金", "GUI 确认后同步关联持仓、交易和批次"],
        ],
    )
    doc.add_paragraph(
        "V1.0 为本地单用户桌面软件，不包含远程账户认证、云端同步或交易账户凭据。"
        "用户应自行保护本地业务数据和备份文件。"
    )

    doc.add_heading("14 第三方依赖与原创边界", level=1)
    doc.add_heading("14.1 第三方依赖", level=2)
    add_table(
        doc,
        ["依赖", "用途", "原创归属"],
        [
            ["xalpha 0.12.3", "获取基金信息、历史净值和公开数据", "第三方"],
            ["pandas 1.5.3", "表格、时间序列和 CSV 处理", "第三方"],
            ["numpy 1.26.4", "数值计算", "第三方"],
            ["matplotlib 3.10.9", "图表绘制和 Tkinter 嵌入", "第三方"],
            ["Tkinter/ttk", "桌面 GUI 控件", "Python 标准库"],
            ["Python 标准库", "JSON、文件、日志、日期、线程、队列等", "标准基础能力"],
        ],
    )
    doc.add_paragraph("上述第三方库及其源码不属于本软件原创代码，不作为原创源程序鉴别材料。")
    doc.add_heading("14.2 原创业务逻辑", level=2)
    add_bullets(
        doc,
        [
            "GUI 页面组织、交互流程、后台分析状态和数据质量展示。",
            "基金池配置、校验、增删改和启停。",
            "xalpha 数据源封装、正式净值补全、缓存与错误回退。",
            "持仓、交易、批次、FIFO 和数据迁移。",
            "统一持仓估值、组合摘要及跨页面一致性。",
            "单基金与组合收益风险指标的组织和计算流程。",
            "固定定投、动态定投和均线策略回测。",
            "策略评分、定投倍率、风险信号和解释生成。",
            "CSV、PNG、Markdown 报告和运行日志输出。",
        ],
    )

    doc.add_heading("15 已知限制", level=1)
    add_bullets(
        doc,
        [
            "第三方数据源可能因网络、接口或页面变化而失败；缓存不代表当日已更新。",
            "QDII 正式净值通常滞后一个或多个交易日；累计净值可能晚于单位净值补全。",
            "组合截止到所有有效成分共同拥有的最新正式净值日。",
            "历史持仓缺少真实份额时只能估算。",
            "组合历史净值按当前持仓权重合成，不还原历史动态权重。",
            "回测未完整模拟费用、滑点、限购、暂停申购、到账延迟和全部交易约束。",
            "策略建议是固定人工规则，不是机器学习或未来价格预测。",
            "卖出采用 FIFO，V1.0 不能直接撤销卖出交易。",
            "GUI 主要适配 Windows；V1.0 尚无独立安装包、完整自动化测试和 CI。",
        ],
    )

    doc.add_heading("16 源码鉴别材料建议范围", level=1)
    doc.add_paragraph(
        "软著源码材料应选择本项目原创业务代码，不包含第三方库源码。建议覆盖以下核心模块："
    )
    add_bullets(
        doc,
        [
            "main.py、app_info.py、ui_app.py、run_portfolio.py。",
            "portfolio_analysis.py、portfolio_store.py、position_service.py。",
            "fund_registry.py、data_provider.py、quant_service.py、app_logging.py。",
            "indicators/、backtest/、recommendation/ 下的原创模块。",
        ],
    )

    doc.add_heading("17 设计结论", level=1)
    doc.add_paragraph(
        "V1.0 已形成从基金配置、真实净值、持仓批次、统一估值到指标、回测、建议和报告的完整本地软件链路。"
        "软件身份、入口、数据目录、输出目录、免责声明、第三方边界和已知限制已统一。"
        "后续申请准备应以冻结基线为准，重点完成标准示例流程验收、测试记录、截图定稿和鉴别材料排版，"
        "不再扩大 V1.0 功能范围。"
    )
    p = doc.add_paragraph(DISCLAIMER)
    for run in p.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)

    path = DOCS / f"{APP_NAME}_软件设计说明书.docx"
    doc.save(path)
    return path


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    user_manual = build_user_manual()
    design_spec = build_design_spec()
    print(user_manual)
    print(design_spec)


if __name__ == "__main__":
    main()

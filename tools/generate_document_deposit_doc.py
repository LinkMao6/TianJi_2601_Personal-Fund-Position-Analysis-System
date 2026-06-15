from __future__ import annotations

import math
import re
import textwrap
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MANUAL = DOCS / "天玑个人基金组合分析与回测系统 V1.0_用户操作手册.docx"
DESIGN = DOCS / "天玑个人基金组合分析与回测系统 V1.0_软件设计说明书.docx"
OUTPUT = DOCS / "天玑个人基金组合分析与回测系统 V1.0_文档鉴别材料_优化版.docx"
SCREENSHOTS = DOCS / "screenshots"

APP_NAME = "天玑个人基金组合分析与回测系统 V1.0"
DOC_NAME = "软件说明书 V1.0"
LINES_PER_PAGE = 30
LINE_WIDTH = 52


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def wrap_text(text: str) -> list[str]:
    text = clean(text)
    if not text:
        return []
    return textwrap.wrap(
        text,
        width=LINE_WIDTH,
        break_long_words=True,
        break_on_hyphens=False,
        replace_whitespace=False,
    )


def iter_blocks(parent):
    if isinstance(parent, DocumentObject):
        parent_element = parent.element.body
    else:
        raise TypeError("Unsupported parent")
    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def extract_document(path: Path, part_title: str) -> list[str]:
    doc = Document(path)
    lines = [part_title]

    for block in iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = clean(block.text)
            if not text:
                continue
            # Remove source-document administration pages. The unified material
            # has one identity, one header and one continuous page sequence.
            if text in {"文档说明", "目录"}:
                continue
            if "请在 Word 中右键更新域" in text:
                continue
            if text == APP_NAME or text in {"用户操作手册", "软件设计说明书"}:
                continue
            if text.startswith("运行、数据管理、") or text.startswith("总体架构、模块划分、"):
                continue
            if block.style.name == "Title":
                continue
            prefix = ""
            if block.style.name.startswith("Heading"):
                prefix = "【"
                text = text + "】"
            elif block.style.name.startswith("List"):
                prefix = "• "
            lines.extend(wrap_text(prefix + text))
        else:
            for row in block.rows:
                values = [clean(cell.text) for cell in row.cells]
                if not any(values):
                    continue
                lines.extend(wrap_text(" | ".join(values)))

    return lines


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(f"{APP_NAME}  {DOC_NAME}")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("第 ")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    add_page_field(footer)
    run = footer.add_run(" 页")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def add_page(
    doc: Document,
    page_lines: list[str],
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(15.4)
    paragraph.paragraph_format.keep_together = True

    for index, text in enumerate(page_lines):
        run = paragraph.add_run(text or " ")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
        if text.startswith("【") or text.startswith("第一篇") or text.startswith("第二篇"):
            run.bold = True
        if index < len(page_lines) - 1:
            run.add_break()


def add_figure_page(
    doc: Document,
    title: str,
    description: list[str],
    image: Path,
    caption: str,
) -> None:
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(8)
    run = heading.add_run(title)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(15)
    run.bold = True

    for text in description:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.35
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)

    picture = doc.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.space_before = Pt(6)
    picture.paragraph_format.space_after = Pt(2)
    picture.add_run().add_picture(str(image), width=Cm(16.0))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(0)
    run = cap.add_run(caption)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


FIGURES = [
    {
        "keyword": "3 主界面与通用操作",
        "title": "主界面与功能导航",
        "image": SCREENSHOTS / "01_总览与图表.png",
        "caption": "图1  总览与图表页面（V1.0模拟演示数据）",
        "description": [
            "软件启动后进入投资组合仪表盘。页眉显示软件全称、版本号、模拟演示数据标识和分析状态，标签栏提供十二个功能页面。",
            "总览页面集中展示组合总资产、累计收益、收益率、盈利与亏损基金数量，并通过资产配置、类别收益、风险收益矩阵和定投对比图辅助观察组合情况。",
        ],
    },
    {
        "keyword": "4 基金池管理",
        "title": "基金池配置与维护",
        "image": SCREENSHOTS / "04_基金池管理.png",
        "caption": "图2  基金池管理页面",
        "description": [
            "基金池管理页面用于添加、编辑、启用、停用和删除基金，并维护基金代码、名称、资产类别、基金类型、数据源和定投配置。",
            "基金代码是持仓、交易、缓存和建议模块之间的关联标识。删除基金前软件会进行确认，并处理相应的关联数据。",
        ],
    },
    {
        "keyword": "5 数据更新与缓存",
        "title": "数据来源、缓存与质量状态",
        "image": SCREENSHOTS / "05_数据中心.png",
        "caption": "图3  数据中心页面",
        "description": [
            "数据中心用于核对每只基金的净值截止日期、数据来源、缓存状态、是否过期和错误信息，是判断本轮分析数据是否可靠的主要页面。",
            "远程更新失败时软件使用最近一次成功保存的真实缓存。QDII等披露滞后基金应以正式净值日期为准，不能仅按当前自然日判断是否更新。",
        ],
    },
    {
        "keyword": "6 持仓、交易与批次",
        "title": "持仓批次与剩余份额管理",
        "image": SCREENSHOTS / "03_持仓批次.png",
        "caption": "图4  持仓批次管理页面",
        "description": [
            "持仓批次页面按基金展示买入日期、确认金额、确认净值、确认份额、剩余份额、成本、手续费和数据质量。",
            "同一基金可以存在多笔买入批次。卖出交易按先进先出规则消耗批次，统一估值依据剩余确认份额和最新正式净值计算。",
        ],
    },
    {
        "keyword": "8 定投模拟与策略回测",
        "title": "策略回测参数与结果",
        "image": SCREENSHOTS / "06_策略回测.png",
        "caption": "图5  策略回测页面",
        "description": [
            "策略回测页面允许选择基金、起止日期、投入频率和基础金额，并比较固定定投、动态定投和均线策略。",
            "结果区域展示累计投入、期末资产、收益金额、收益率、最大回撤等指标，下方资金曲线用于观察策略资产随时间的变化。",
        ],
    },
    {
        "keyword": "9 策略建议",
        "title": "策略建议生成与风险提示",
        "image": SCREENSHOTS / "07_策略建议.png",
        "caption": "图6  策略建议页面",
        "description": [
            "策略建议页面根据量化指标、持仓仓位、基金类型和定投配置生成规则化建议，可选择全部基金或仅当前持仓基金。",
            "建议结果包括动作、评分、建议金额、置信度、数据质量、生成理由和风险信号。建议仅供个人学习与研究参考，不构成投资建议。",
        ],
    },
]


def paginate(lines: list[str]) -> list[list[str]]:
    return [lines[i : i + LINES_PER_PAGE] for i in range(0, len(lines), LINES_PER_PAGE)]


def main() -> None:
    manual_lines = extract_document(MANUAL, "第一篇 用户操作说明")
    design_lines = extract_document(DESIGN, "第二篇 软件设计说明")

    # The unified documentation is shorter than 60 pages, so the applicable
    # ordinary-deposit form is the complete document, not padded or duplicated
    # material. Full pages contain 30 lines; only the natural final page may
    # contain fewer.
    manual_pages = paginate(manual_lines)
    design_pages = paginate(design_lines)
    total_pages = len(manual_pages) + len(design_pages)

    doc = Document()
    configure(doc)
    all_pages = manual_pages + design_pages
    inserted = set()
    for page_index, page_lines in enumerate(all_pages):
        add_page(doc, page_lines)
        matches = [
                figure
                for figure in FIGURES
                if any(
                    line.startswith("【") and figure["keyword"] in line
                    for line in page_lines
                )
                and figure["keyword"] not in inserted
        ]
        for matching in matches:
            doc.add_page_break()
            add_figure_page(
                doc,
                matching["title"],
                matching["description"],
                matching["image"],
                matching["caption"],
            )
            inserted.add(matching["keyword"])
        if page_index < total_pages - 1:
            doc.add_page_break()

    props = doc.core_properties
    props.title = f"{APP_NAME} {DOC_NAME}"
    props.subject = "计算机软件著作权登记文档鉴别材料"
    props.comments = (
        "统一软件说明书，由用户操作说明和软件设计说明两个独立起页的篇章组成；"
        "A4版式，提交全文；普通完整正文页30行，图示页紧邻对应功能说明并配有文字解释，"
        "各篇独立起页、末页自然结束，页码连续。"
    )

    doc.save(OUTPUT)
    print(f"output={OUTPUT}")
    print(f"manual_available={len(manual_lines)}")
    print(f"design_available={len(design_lines)}")
    print(f"selected={len(manual_lines) + len(design_lines)}")
    print(f"pages={total_pages}")
    print(f"manual_pages={len(manual_pages)}")
    print(f"design_pages={len(design_pages)}")
    print(f"figures={len(inserted)}")


if __name__ == "__main__":
    main()
